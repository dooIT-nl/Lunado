#!/usr/bin/env python3
"""Generate the d1_shipping_cost user manual in Word (.docx) format.

Houd dit script gelijk met de module. Bij een functionele wijziging:
1. Werk de betreffende paragraaf/tabel hieronder bij.
2. Verhoog MANUAL_VERSION / MANUAL_DATE.
3. Voeg een regel toe aan "Bijlage C — Wijzigingshistorie".
4. Draai het script opnieuw om de .docx te regenereren.
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE

MANUAL_VERSION = "1.9"
MANUAL_DATE = "Juli 2026"

doc = Document()

# -- Styles --
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

for level in range(1, 4):
    hs = doc.styles[f"Heading {level}"]
    hs.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)  # dark green


def add_table(headers, rows):
    """Add a formatted table to the document."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.rows[ri + 1].cells[ci].text = str(val)
    doc.add_paragraph()  # spacing


# ============================================================
# TITLE PAGE
# ============================================================
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Gebruikershandleiding\nTransportkosten (d1_shipping_cost)")
run.font.size = Pt(26)
run.bold = True
run.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Odoo 19 — dooIT B.V.")
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()
doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run(f"Versie {MANUAL_VERSION} — {MANUAL_DATE}").font.size = Pt(11)

doc.add_page_break()

# ============================================================
# TABLE OF CONTENTS (manual)
# ============================================================
doc.add_heading("Inhoudsopgave", level=1)
toc_items = [
    "1. Inleiding",
    "2. Installatie en configuratie",
    "   2.1 Module installeren",
    "   2.2 Transportklassen",
    "   2.3 Lengteklassen",
    "   2.4 Transporttarieven",
    "   2.5 Postcode-prefixen en adresfiltering",
    "   2.6 Instellingen (drempelwaarden)",
    "3. Producten configureren",
    "   3.1 Transportvelden",
    "   3.2 Hoeveelheid × lengte (staaf-/lengteproducten)",
    "4. Orderregels: Hoeveelheid en Lengte",
    "5. Transportkosten berekenen",
    "   5.1 Berekening starten",
    "   5.2 Vaste verzendwijze klant (berekening overslaan)",
    "   5.3 Resultaat, transporteur en transportkostenregel",
    "   5.4 Palletberekening",
    "   5.5 Herberekening",
    "6. Lijstweergave en filteren",
    "",
    "Bijlage A — Berekeningsflows TK1 / TK2 / TK3",
    "Bijlage B — Transporteurbepaling en gewicht per stuk",
    "Bijlage C — Open punten",
    "Bijlage D — Wijzigingshistorie",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ============================================================
# 1. INLEIDING
# ============================================================
doc.add_heading("1. Inleiding", level=1)
doc.add_paragraph(
    "De module Transportkosten (d1_shipping_cost) berekent automatisch de "
    "verzendkosten op offertes en verkooporders in Odoo. De berekening is "
    "gebaseerd op drie transportklassen:"
)
add_table(
    ["Klasse", "Type", "Eenheid", "Voorbeeld"],
    [
        ["TK1", "Bundels (lang)", "Bundels", "Buizen, profielen"],
        ["TK2", "Dozen (gewicht)", "Dozen", "Koppelingen, fittings"],
        ["TK3", "Dozen (display)", "Dozen", "Displays, borden"],
    ],
)
doc.add_paragraph(
    "Per klasse wordt een apart deelproces doorlopen. De som van TK1 + TK2 + TK3 "
    "wordt als één transportkostenregel aan de order toegevoegd. Op basis van de "
    "gevonden tarieven wordt bovendien automatisch de transporteur (verzendwijze) "
    "op de order bepaald. Wanneer de kosten niet automatisch bepaald kunnen worden "
    "(bv. te veel bundels, te zwaar of een uitzonderingsmaat), wordt de order "
    "gemarkeerd voor handmatige palletberekening."
)
doc.add_paragraph(
    "Heeft de klant zelf een vaste verzendwijze afgesproken, dan wordt de "
    "automatische berekening overgeslagen (zie §5.2)."
)

# ============================================================
# 2. INSTALLATIE EN CONFIGURATIE
# ============================================================
doc.add_heading("2. Installatie en configuratie", level=1)

doc.add_heading("2.1 Module installeren", level=2)
doc.add_paragraph(
    "De module wordt geïnstalleerd via de Odoo Apps-lijst (zoek op "
    "\"Transportkosten\") of via de command line:"
)
doc.add_paragraph("odoo-bin -i d1_shipping_cost --stop-after-init --no-http", style="No Spacing")
doc.add_paragraph(
    "De module is afhankelijk van de standaardmodules Verkoop (sale_management), "
    "Levering (delivery) en Producten (product). De koppeling met Levering is nodig "
    "voor de transporteurs (delivery.carrier) en de vaste verzendwijze per klant."
)
doc.add_paragraph()

doc.add_heading("2.2 Transportklassen", level=2)
doc.add_paragraph("Ga naar: Verkoop → Configuratie → Transport → Transportklassen")
doc.add_paragraph(
    "Hier staan de transportklassen gedefinieerd. Standaard worden TK1, TK2 en TK3 "
    "aangemaakt. Elke klasse heeft een code (hoofdletters, bv. TK1) die intern "
    "gebruikt wordt om de juiste berekeningsflow te selecteren."
)
add_table(
    ["Veld", "Omschrijving"],
    [
        ["Naam", "Weergavenaam, bv. 'TK1 — Bundels'"],
        ["Code", "Technische code (TK1, TK2, TK3). Hoofdletters."],
        ["Omschrijving", "Vrije toelichting"],
        ["Volgorde", "Sorteervolgorde in keuzelijsten"],
    ],
)

doc.add_heading("2.3 Lengteklassen", level=2)
doc.add_paragraph("Ga naar: Verkoop → Configuratie → Transport → Lengteklassen")
doc.add_paragraph(
    "Lengteklassen definiëren bereiken in centimeters voor de TK1-tariefopzoeking. "
    "De maximale productlengte in een order bepaalt welke lengteklasse van toepassing is."
)
add_table(
    ["Veld", "Omschrijving"],
    [
        ["Naam", "Weergavenaam, bv. '< 1,65 m'"],
        ["Lengte vanaf (cm)", "Ondergrens van het bereik (inclusief)"],
        ["Lengte t/m (cm)", "Bovengrens (inclusief). 0 = open einde (geen bovengrens)."],
    ],
)
doc.add_paragraph("Standaard lengteklassen:")
add_table(
    ["Naam", "Van (cm)", "T/m (cm)"],
    [
        ["< 1,65 m", "0", "164"],
        ["1,65 – 2,14 m", "165", "214"],
        [">= 2,15 m", "215", "0 (open)"],
    ],
)
doc.add_paragraph(
    "Tip: Wijzigt de vervoerder de laadruimte-afmetingen? Pas gewoon de cm-waarden "
    "aan of voeg een nieuwe lengteklasse toe. Er is geen codewijziging nodig."
)

doc.add_heading("2.4 Transporttarieven", level=2)
doc.add_paragraph("Ga naar: Verkoop → Configuratie → Transport → Transporttarieven")
doc.add_paragraph(
    "Hier worden de tariefstaffels beheerd. Elk tarief koppelt een transportklasse, "
    "optioneel een lengteklasse, een eenheidsstaffel (van/t.m.), een adresfilter en "
    "een transporteur aan een prijs."
)
add_table(
    ["Veld", "Omschrijving"],
    [
        ["Naam", "Herkenbare naam voor het tarief"],
        ["Transportklasse", "TK1, TK2 of TK3"],
        ["Lengteklasse", "Alleen voor TK1. Laat leeg voor TK2/TK3."],
        ["Eenheidstype", "Bundels (TK1) of Dozen (TK2/TK3)"],
        ["Aantal vanaf / t.m.", "Staffelbereik. 0 als bovengrens = open einde."],
        ["Tarief", "Prijs voor deze staffel"],
        ["Transporteur", "De verzendwijze (delivery.carrier) die bij dit tarief hoort. "
                         "Bepaalt de transporteur op de order (zie Bijlage B)."],
        ["Landen", "Adresfilter: alleen voor deze landen (leeg = alle)"],
        ["Provincies", "Adresfilter: alleen voor deze provincies"],
        ["Postcode-prefixen", "Adresfilter: postcode moet beginnen met prefix"],
    ],
)

doc.add_heading("2.5 Postcode-prefixen en adresfiltering", level=2)
doc.add_paragraph(
    "De adresfiltering werkt identiek aan de standaard Odoo delivery.carrier module. "
    "Per tariefregel kun je optioneel instellen:"
)
doc.add_paragraph("• Landen — tarief geldt alleen voor deze landen", style="List Bullet")
doc.add_paragraph("• Provincies — gefilterd op geselecteerde landen", style="List Bullet")
doc.add_paragraph(
    "• Postcode-prefixen — tarief geldt alleen als de postcode van het afleveradres "
    "begint met een van de opgegeven prefixen",
    style="List Bullet",
)
doc.add_paragraph(
    "Als alle drie leeg zijn, geldt het tarief als generiek (voor alle adressen). "
    "Reguliere expressies worden ondersteund in postcode-prefixen. Bij meerdere "
    "passende tarieven wint het meest specifieke adresfilter (postcode > provincie > "
    "land > generiek)."
)
doc.add_paragraph("Voorbeelden:")
add_table(
    ["Prefix", "Matcht", "Matcht niet"],
    [
        ["8861", "8861AA, 8861 ZZ", "8862AA"],
        ["886", "8861AA, 8862BB, 8869ZZ", "8870AA"],
        ["8861$", "alleen exact '8861'", "8861AA"],
    ],
)

doc.add_heading("2.6 Instellingen (drempelwaarden)", level=2)
doc.add_paragraph("Ga naar: Verkoop → Instellingen → sectie Transportkosten")
doc.add_paragraph("Hier stel je de drempelwaarden en afmetingen in:")
add_table(
    ["Parameter", "Standaard", "Toelichting"],
    [
        ["TK1 max bundels", "99999", "Vanaf dit aantal bundels → palletberekening"],
        ["Bundelbreedte (cm)", "15", "Breedte van een standaardbundel"],
        ["Bundelhoogte (cm)", "15", "Hoogte van een standaardbundel"],
        ["Max gewicht bundel (kg)", "20", "Max gewicht per bundel (ook gebruikt voor gewichtsbanding)"],
        ["TK2 max gewicht (kg)", "99999", "Vanaf dit totaalgewicht → palletberekening"],
        ["TK3 max gewicht (kg)", "99999", "Vanaf dit totaalgewicht → palletberekening"],
        ["Max gewicht doos (kg)", "20", "Max gewicht per doos (TK2/TK3, ook voor gewichtsbanding)"],
    ],
)

doc.add_page_break()

# ============================================================
# 3. PRODUCTEN CONFIGUREREN
# ============================================================
doc.add_heading("3. Producten configureren", level=1)

doc.add_heading("3.1 Transportvelden", level=2)
doc.add_paragraph("Ga naar een product → tabblad Transport")
doc.add_paragraph("Stel de volgende velden in:")
add_table(
    ["Veld", "Omschrijving"],
    [
        ["Transportklasse", "Kies TK1, TK2 of TK3. Producten zonder klasse worden overgeslagen."],
        ["Lengte (cm)", "Productlengte in centimeters"],
        ["Breedte (cm)", "Productbreedte in centimeters"],
        ["Hoogte (cm)", "Producthoogte in centimeters"],
        ["Gebruik Hoeveelheid", "Aan = de orderregel-hoeveelheid wordt berekend als "
                                 "Hoeveelheid × Lengte (zie §3.2 en §4)."],
        ["Gebruik Lengte", "Aan = de gebruiker mag op de orderregel een eigen lengte "
                            "invoeren. Uit = de productlengte wordt gebruikt."],
    ],
)
doc.add_paragraph(
    "Het gewicht wordt gelezen uit het standaard Odoo-veld Gewicht op het product. "
    "Dit hoeft niet apart ingesteld te worden op het Transport-tabblad."
)

doc.add_heading("3.2 Hoeveelheid × lengte (staaf-/lengteproducten)", level=2)
doc.add_paragraph(
    "Voor producten die per lengte-eenheid worden verkocht (bijvoorbeeld per meter) "
    "kan het handig zijn om op de orderregel het aantal stuks én de lengte per stuk "
    "apart in te voeren. Zet daarvoor op het product Gebruik Hoeveelheid aan."
)
doc.add_paragraph("Gedrag:")
doc.add_paragraph(
    "• Gebruik Hoeveelheid aan, Gebruik Lengte uit → de lengte wordt automatisch "
    "overgenomen van de productlengte (in meters). De besteleenheid (Aantal) wordt "
    "berekend als Hoeveelheid × Lengte.",
    style="List Bullet",
)
doc.add_paragraph(
    "• Gebruik Hoeveelheid aan, Gebruik Lengte aan → de gebruiker voert zelf een "
    "lengte per stuk in op de orderregel. Aantal = Hoeveelheid × Lengte.",
    style="List Bullet",
)
doc.add_paragraph(
    "• Gebruik Hoeveelheid uit → normaal gedrag: de gebruiker voert het Aantal "
    "handmatig in.",
    style="List Bullet",
)
doc.add_paragraph(
    "Bij de gewichtsberekening houdt de module rekening met de maateenheid van het "
    "product: is het product een lengte-eenheid (m, cm, mm …), dan wordt het "
    "productgewicht (per eenheid) vermenigvuldigd met de lengte om het werkelijke "
    "stukgewicht te bepalen (zie Bijlage B)."
)

# ============================================================
# 4. ORDERREGELS: HOEVEELHEID EN LENGTE
# ============================================================
doc.add_heading("4. Orderregels: Hoeveelheid en Lengte", level=1)
doc.add_paragraph(
    "Op de orderregels van een offerte/verkooporder zijn twee extra kolommen "
    "beschikbaar (standaard zichtbaar, in te klappen via het kolommenmenu):"
)
add_table(
    ["Kolom", "Omschrijving"],
    [
        ["Hoeveelheid", "Aantal stuks. Alleen invoerbaar bij producten met "
                        "'Gebruik Hoeveelheid' aan."],
        ["Lengte", "Lengte per stuk (in meters). Alleen invoerbaar bij producten met "
                   "'Gebruik Lengte' aan; anders automatisch gevuld vanuit het product."],
    ],
)
doc.add_paragraph(
    "Wanneer 'Gebruik Hoeveelheid' actief is, wordt het standaard Odoo-veld Aantal "
    "automatisch berekend als Hoeveelheid × Lengte. Deze berekening werkt zowel in "
    "de UI (bij het wijzigen van velden) als bij het aanmaken/wijzigen van regels via "
    "import of API."
)

# ============================================================
# 5. TRANSPORTKOSTEN BEREKENEN
# ============================================================
doc.add_heading("5. Transportkosten berekenen", level=1)

doc.add_heading("5.1 Berekening starten", level=2)
doc.add_paragraph(
    "Open een offerte of verkooporder en klik op de knop \"Bereken transportkosten\" "
    "(🚛-icoon) in de header. De module doorloopt per aanwezige transportklasse "
    "de bijbehorende berekeningsflow en voegt één getotaliseerde transportkostenregel "
    "toe aan de order."
)
doc.add_paragraph(
    "Belangrijk: alleen orderregels met een product waarop een transportklasse is "
    "ingesteld worden meegenomen. Regels zonder klasse worden genegeerd."
)

doc.add_heading("5.2 Vaste verzendwijze klant (berekening overslaan)", level=2)
doc.add_paragraph(
    "Heeft de klant een vaste verzendwijze afgesproken, dan hoeft de module niets te "
    "berekenen. Dit wordt bepaald via het standaardveld Verzendwijze "
    "(property_delivery_carrier_id) op de klant (tabblad Verkoop & Aankoop van de "
    "contactpersoon)."
)
doc.add_paragraph(
    "Als dit veld gevuld is en je klikt op \"Bereken transportkosten\", dan wordt de "
    "TK1/TK2/TK3-berekening volledig overgeslagen. Er wordt géén transportkostenregel "
    "toegevoegd en er verschijnt een bericht in de Chatter, bijvoorbeeld:"
)
p = doc.add_paragraph(style="No Spacing")
p.add_run("Transportkostenberekening").bold = True
doc.add_paragraph(
    "Verzendwijze klant: <naam verzendwijze>, transportkostenberekening overgeslagen",
    style="No Spacing",
)
doc.add_paragraph()
doc.add_paragraph(
    "Alleen wanneer de klant géén vaste verzendwijze heeft, wordt de automatische "
    "berekening uitgevoerd."
)

doc.add_heading("5.3 Resultaat, transporteur en transportkostenregel", level=2)
doc.add_paragraph("Na een uitgevoerde berekening verschijnt onder de orderregels:")
doc.add_paragraph(
    "• Een samenvattingsveld (Transport berekening) met per klasse het resultaat "
    "(aantal bundels/dozen, tarief, of reden voor handmatige berekening).",
    style="List Bullet",
)
doc.add_paragraph(
    "• De bepaalde transporteur. De module kiest de transporteur op basis van de "
    "gevonden tarieven, met prioriteit TK1 > TK3 > TK2 (zie Bijlage B). De "
    "transporteur wordt op de order gezet (veld Verzendwijze / carrier_id).",
    style="List Bullet",
)
doc.add_paragraph(
    "• Eén transportkostenregel onderaan de order. Als product wordt het product van "
    "de gekozen transporteur gebruikt; de regel krijgt de markering [TRANSPORT] en het "
    "totaalbedrag als prijs.",
    style="List Bullet",
)
doc.add_paragraph(
    "• Een chatterbericht met dezelfde samenvatting (voor audit trail).",
    style="List Bullet",
)
doc.add_paragraph(
    "Let op: is er geen transporteur (of geen product op de transporteur) bepaald "
    "terwijl er wel kosten zijn, dan geeft de module een foutmelding en wordt er geen "
    "regel aangemaakt. Controleer in dat geval de tarieven en de transporteur-"
    "configuratie."
)

doc.add_heading("5.4 Palletberekening", level=2)
doc.add_paragraph(
    "Het veld Palletberekening (checkbox) wordt aangevinkt wanneer de kosten niet "
    "automatisch bepaald konden worden — bijvoorbeeld bij te veel bundels, te zwaar, "
    "een ontbrekend tarief of een TK3-uitzonderingsmaat."
)
doc.add_paragraph(
    "Wanneer Palletberekening is aangevinkt, wordt er géén (of slechts een gedeeltelijke) "
    "automatische transportkostenregel toegevoegd voor het handmatige deel. De order "
    "moet dan handmatig beoordeeld worden. De samenvatting vermeldt precies welk "
    "onderdeel handmatige behandeling vereist."
)

doc.add_heading("5.5 Herberekening", level=2)
doc.add_paragraph(
    "Bij het opnieuw klikken op \"Bereken transportkosten\" wordt de bestaande "
    "transportkostenregel (herkend aan de [TRANSPORT]-markering) automatisch "
    "verwijderd en vervangen door een nieuwe berekening. Er worden dus nooit dubbele "
    "transportkostenregels aangemaakt."
)

# ============================================================
# 6. LIJSTWEERGAVE EN FILTEREN
# ============================================================
doc.add_heading("6. Lijstweergave en filteren", level=1)
doc.add_paragraph(
    "In de lijst van verkooporders is de kolom Palletberekening zichtbaar "
    "(na de klantkolom). Daarnaast is er een zoekfilter \"Palletberekening\" "
    "beschikbaar om snel alle orders te vinden die handmatige beoordeling vereisen."
)

doc.add_page_break()

# ============================================================
# BIJLAGE A — BEREKENINGSFLOWS
# ============================================================
doc.add_heading("Bijlage A — Berekeningsflows TK1 / TK2 / TK3", level=1)

# --- TK1 ---
doc.add_heading("A.1 TK1 — Bundelberekening (lange producten)", level=2)
doc.add_paragraph(
    "Invoer: alle orderregels met transportklasse TK1. Per product zijn breedte, "
    "hoogte, lengte (cm) en gewicht (kg) bekend. Het aantal stuks komt uit het veld "
    "Hoeveelheid (bij 'Gebruik Hoeveelheid') of anders uit Aantal."
)

doc.add_heading("Stap 1) Gewichtsbanding per stuk", level=3)
doc.add_paragraph(
    "Vóór de bundelberekening worden de stuks op basis van het stukgewicht in drie "
    "banden verdeeld (max = Max gewicht bundel, standaard 20 kg):"
)
doc.add_paragraph(
    "• Stukgewicht >= max → de volledige TK1-klasse gaat naar PALLETBEREKENING.",
    style="List Bullet",
)
doc.add_paragraph(
    "• max/2 <= stukgewicht < max → 'één-per-bundel'-pool: elk stuk vormt een eigen "
    "bundel.",
    style="List Bullet",
)
doc.add_paragraph(
    "• Stukgewicht < max/2 → normale (geometrische) pool.",
    style="List Bullet",
)
doc.add_paragraph(
    "Het stukgewicht houdt rekening met de maateenheid: bij lengteproducten is dat "
    "productgewicht × lengte (zie Bijlage B)."
)

doc.add_heading("Stap 2) Buizen per bundel (normale pool)", level=3)
doc.add_paragraph(
    "Per product in de normale pool wordt berekend hoeveel stuks in de dwarsdoorsnede "
    "van een standaardbundel passen:"
)
p = doc.add_paragraph()
p.add_run("buizen_in_breedte").bold = True
p.add_run(" = floor(bundelbreedte / artikelbreedte)")
p = doc.add_paragraph()
p.add_run("buizen_in_hoogte").bold = True
p.add_run(" = floor(bundelhoogte / artikelhoogte)")
p = doc.add_paragraph()
p.add_run("buizen_per_bundel").bold = True
p.add_run(" = buizen_in_breedte × buizen_in_hoogte")
doc.add_paragraph(
    "floor = naar beneden afronden. Als de artikelafmeting 0 of groter dan de "
    "bundelmaat is, wordt minimaal 1 aangehouden."
)
p = doc.add_paragraph()
p.add_run("bundels_normaal").bold = True
p.add_run(" = ceil(aantal_stuks / buizen_per_bundel), per product opgeteld")

doc.add_heading("Stap 3) Gewichtscorrectie (normale pool)", level=3)
doc.add_paragraph("Controleer of het gemiddelde gewicht per bundel binnen de limiet valt:")
doc.add_paragraph(
    "• Als (gewicht_normale_pool / bundels_normaal) < max_gewicht_bundel → "
    "bundels_normaal blijft ongewijzigd.",
    style="List Bullet",
)
doc.add_paragraph(
    "• Anders → bundels_normaal = ceil(gewicht_normale_pool / max_gewicht_bundel).",
    style="List Bullet",
)

doc.add_heading("Stap 4) Totaal bundels", level=3)
p = doc.add_paragraph()
p.add_run("totaal_bundels").bold = True
p.add_run(" = bundels_één-per-bundel + bundels_normaal")

doc.add_heading("Stap 5) Drempel- en tariefbeslissing", level=3)
doc.add_paragraph(
    "• Als totaal_bundels >= TK1 max bundels (instelbaar) → PALLETBEREKENING "
    "(handmatig). Geen automatische prijs.",
    style="List Bullet",
)
doc.add_paragraph(
    "• Anders → bepaal de lengteklasse op basis van de maximale productlengte in de "
    "order (opzoeking in de tabel Lengteklassen) en zoek het tarief op in de "
    "Transporttarieven (klasse TK1 + lengteklasse + staffel + adresfilter). Uit het "
    "tarief volgt tevens de transporteur.",
    style="List Bullet",
)

doc.add_heading("Rekenvoorbeeld TK1", level=3)
add_table(
    ["Gegeven", "Waarde"],
    [
        ["Product", "Aluminium buis 60×3 mm, lengte 150 cm, gewicht 2 kg/stuk"],
        ["Bundelafmeting", "15×15 cm (standaard)"],
        ["Max gewicht bundel", "20 kg"],
        ["Besteld aantal", "50 stuks"],
    ],
)
doc.add_paragraph("Berekening:")
doc.add_paragraph("1. stukgewicht 2 kg < 10 (= 20/2) → normale pool", style="List Number")
doc.add_paragraph("2. buizen_in_breedte = floor(15 / 3.0) = 5", style="List Number")
doc.add_paragraph("3. buizen_in_hoogte = floor(15 / 3.0) = 5", style="List Number")
doc.add_paragraph("4. buizen_per_bundel = 5 × 5 = 25", style="List Number")
doc.add_paragraph("5. bundels_normaal = ceil(50 / 25) = 2", style="List Number")
doc.add_paragraph("6. gewicht_normale_pool = 50 × 2 = 100 kg", style="List Number")
doc.add_paragraph("7. gemiddeld gewicht/bundel = 100 / 2 = 50 kg > 20 kg → correctie!", style="List Number")
doc.add_paragraph("8. bundels_normaal = ceil(100 / 20) = 5 → totaal_bundels = 5", style="List Number")
doc.add_paragraph("9. max lengte = 150 cm → lengteklasse '< 1,65 m'", style="List Number")
doc.add_paragraph("10. Tariefopzoeking: TK1, < 1,65 m, 5 bundels → € 25,00 (+ transporteur)", style="List Number")

doc.add_paragraph()

# --- TK2 ---
doc.add_heading("A.2 TK2 — Doosberekening (op gewicht)", level=2)
doc.add_paragraph("Invoer: alle orderregels met transportklasse TK2.")

doc.add_heading("Stap 1) Totaalgewicht-drempel", level=3)
doc.add_paragraph(
    "• Als totaal_gewicht >= TK2 max gewicht (instelbaar) → PALLETBEREKENING.",
    style="List Bullet",
)

doc.add_heading("Stap 2) Gewichtsbanding per stuk", level=3)
doc.add_paragraph(
    "Per stuk wordt op basis van het stukgewicht ingedeeld (max = Max gewicht doos, "
    "standaard 20 kg):"
)
doc.add_paragraph(
    "• Stukgewicht >= max → volledige TK2-klasse → PALLETBEREKENING.",
    style="List Bullet",
)
doc.add_paragraph(
    "• max/2 <= stukgewicht < max → 'één-per-doos'-pool: elk stuk in een eigen doos.",
    style="List Bullet",
)
doc.add_paragraph(
    "• Stukgewicht < max/2 → lichte pool: gewichten worden opgeteld.",
    style="List Bullet",
)

doc.add_heading("Stap 3) Doosberekening en totaal", level=3)
p = doc.add_paragraph()
p.add_run("dozen_normaal").bold = True
p.add_run(" = ceil(gewicht_lichte_pool / max_gewicht_doos)")
p = doc.add_paragraph()
p.add_run("totaal_dozen").bold = True
p.add_run(" = dozen_één-per-doos + dozen_normaal")

doc.add_heading("Stap 4) Tariefopzoeking", level=3)
doc.add_paragraph(
    "Zoek in Transporttarieven: klasse TK2, staffel op totaal_dozen, adresfilter. "
    "Uit het tarief volgt tevens de transporteur."
)

doc.add_heading("Rekenvoorbeeld TK2", level=3)
add_table(
    ["Gegeven", "Waarde"],
    [
        ["Product", "Koppeling, gewicht 0,5 kg/stuk"],
        ["Max gewicht doos", "20 kg"],
        ["Besteld aantal", "30 stuks"],
    ],
)
doc.add_paragraph("Berekening:")
doc.add_paragraph("1. totaal_gewicht = 30 × 0,5 = 15 kg (< drempel → door)", style="List Number")
doc.add_paragraph("2. stukgewicht 0,5 kg < 10 → lichte pool", style="List Number")
doc.add_paragraph("3. dozen_normaal = ceil(15 / 20) = 1 → totaal_dozen = 1", style="List Number")
doc.add_paragraph("4. Tariefopzoeking: TK2, 1 doos → € 15,00 (+ transporteur)", style="List Number")

doc.add_paragraph()

# --- TK3 ---
doc.add_heading("A.3 TK3 — Displayberekening (volumineus/zwaar)", level=2)
doc.add_paragraph(
    "Invoer: alle orderregels met transportklasse TK3. De beslissingsvolgorde "
    "is strikt:"
)

doc.add_heading("Stap 1) Totaalgewicht-drempel", level=3)
doc.add_paragraph(
    "• Als totaal_gewicht >= TK3 max gewicht → PALLETBEREKENING.",
    style="List Bullet",
)

doc.add_heading("Stap 2) Per-artikel checks", level=3)
doc.add_paragraph("Voor elk product in de TK3-regels wordt gecontroleerd:")
doc.add_paragraph(
    "a) Gewicht per artikel >= 20 kg → PALLETBEREKENING",
    style="List Bullet",
)
doc.add_paragraph(
    "b) Grootste afmeting (L, B of H) >= 165 cm → UITZONDERING Wesseling/Mainfreight "
    "(handmatige behandeling, berekening nog te bepalen)",
    style="List Bullet",
)
doc.add_paragraph(
    "c) Omtrekmaat (L + 2×B + 2×H) >= 300 cm → UITZONDERING Wesseling/Mainfreight",
    style="List Bullet",
)

doc.add_heading("Stap 3) Doosberekening", level=3)
doc.add_paragraph("Als alle checks doorstaan zijn:")
p = doc.add_paragraph()
p.add_run("aantal_dozen").bold = True
p.add_run(" = ceil(totaal_gewicht / max_gewicht_doos)")
doc.add_paragraph(
    "Tariefopzoeking: klasse TK3, staffel op aantal_dozen, adresfilter. Uit het "
    "tarief volgt tevens de transporteur."
)

doc.add_heading("Rekenvoorbeeld TK3", level=3)
add_table(
    ["Gegeven", "Waarde"],
    [
        ["Product", "Display 60×40×30 cm, gewicht 2 kg"],
        ["Besteld aantal", "5 stuks"],
        ["Max gewicht doos", "20 kg"],
    ],
)
doc.add_paragraph("Berekening:")
doc.add_paragraph("1. totaal_gewicht = 5 × 2 = 10 kg (< drempel → door)", style="List Number")
doc.add_paragraph("2. Per artikel: gewicht 2 kg < 20 → OK", style="List Number")
doc.add_paragraph("3. Grootste afmeting = 60 cm < 165 → OK", style="List Number")
doc.add_paragraph("4. Omtrekmaat = 60 + 2×40 + 2×30 = 200 cm < 300 → OK", style="List Number")
doc.add_paragraph("5. aantal_dozen = ceil(10 / 20) = 1 doos", style="List Number")
doc.add_paragraph("6. Tariefopzoeking: TK3, 1 doos → € 20,00 (+ transporteur)", style="List Number")

doc.add_page_break()

# ============================================================
# BIJLAGE B — TRANSPORTEURBEPALING EN GEWICHT PER STUK
# ============================================================
doc.add_heading("Bijlage B — Transporteurbepaling en gewicht per stuk", level=1)

doc.add_heading("B.1 Transporteurbepaling", level=2)
doc.add_paragraph(
    "Elk transporttarief kan een transporteur (delivery.carrier) hebben. Tijdens de "
    "berekening levert elke klasse die een tarief vindt ook een transporteur op. De "
    "module kiest vervolgens één transporteur voor de hele order op basis van de "
    "volgende prioriteit:"
)
p = doc.add_paragraph()
p.add_run("TK1 > TK3 > TK2").bold = True
doc.add_paragraph(
    "De eerste klasse (in deze volgorde) met een transporteur bepaalt de order-"
    "transporteur. De gekozen transporteur wordt op de order gezet (veld "
    "Verzendwijze / carrier_id) en het product van die transporteur wordt gebruikt "
    "voor de transportkostenregel."
)
doc.add_paragraph(
    "Wordt er geen transporteur gevonden terwijl er wel klassen aanwezig zijn, dan "
    "vermeldt de samenvatting 'handmatig selecteren'."
)

doc.add_heading("B.2 Gewicht per stuk bij lengteproducten", level=2)
doc.add_paragraph(
    "Het standaard Odoo-veld Gewicht is het gewicht per maateenheid van het product. "
    "Bij producten die per lengte worden verkocht (maateenheid m, cm, mm …) is dat "
    "dus het gewicht per meter, niet per stuk."
)
doc.add_paragraph(
    "De module detecteert dit via de maateenheid van het product. Als het product een "
    "lengte-eenheid gebruikt én de orderregel een lengte heeft, wordt het stukgewicht "
    "berekend als:"
)
p = doc.add_paragraph()
p.add_run("stukgewicht").bold = True
p.add_run(" = gewicht_per_eenheid × lengte_in_producteenheid")
doc.add_paragraph(
    "Voorbeeld: een buis weegt 3 kg/m en is 2 m lang → stukgewicht = 6 kg. Dit "
    "stukgewicht wordt gebruikt in de gewichtsbanding (TK1/TK2) en de per-artikel "
    "checks (TK3). Voor producten die niet per lengte worden verkocht, blijft het "
    "gewicht simpelweg het productgewicht."
)

doc.add_page_break()

# ============================================================
# BIJLAGE C — OPEN PUNTEN
# ============================================================
doc.add_heading("Bijlage C — Open punten", level=1)
doc.add_paragraph(
    "De volgende punten zijn nog niet definitief vastgesteld en moeten met de "
    "business worden afgestemd:"
)
add_table(
    ["#", "Open punt", "Huidige situatie"],
    [
        ["1", "Drempelwaarde TK1 max bundels",
         "Standaard 99999 (= altijd doorrekenen). Vast te stellen met business."],
        ["2", "Drempelwaarde TK2/TK3 max gewicht (kg)",
         "Standaard 99999. Vast te stellen met business."],
        ["3", "Bundel-/doosafmetingen (15×15 cm, 20 kg)",
         "Zijn aannames. Laten bevestigen door business."],
        ["4", "TK3 uitzondering Wesseling/Mainfreight",
         "Order wordt gevlagd voor handmatige behandeling. "
         "Daadwerkelijke prijsberekening nog te bepalen."],
        ["5", "Tarief-interpretatie",
         "Tarieven worden als totaalbedrag per staffel geïnterpreteerd "
         "(niet per eenheid). Bevestigen met business."],
    ],
)

# ============================================================
# BIJLAGE D — WIJZIGINGSHISTORIE
# ============================================================
doc.add_heading("Bijlage D — Wijzigingshistorie", level=1)
add_table(
    ["Versie", "Wijziging"],
    [
        ["1.9", "Berekening wordt overgeslagen als de klant een vaste verzendwijze "
                "heeft (property_delivery_carrier_id); melding in de Chatter."],
        ["1.8", "Maateenheid-bewust stukgewicht — detectie van lengte-eenheden "
                "(m, cm, mm …) via de product-maateenheid voor TK1/TK2/TK3."],
        ["1.7", "Correctie stukgewicht TK1 voor per-meter verkochte producten "
                "(gewicht/m × lengte = werkelijk stukgewicht)."],
        ["1.6", "Product van de transporteur wordt gebruikt voor de transportkosten-"
                "regel; correctie m→cm in TK1."],
        ["1.5", "i18n — labels Hoeveelheid/Aantal (nl_NL)."],
        ["1.4", "i18n — Engelse bronteksten, Nederlandse (nl_NL) vertalingen."],
        ["1.3", "Transporteurbepaling (prioriteit TK1 > TK3 > TK2); Lengte/Hoeveelheid "
                "op orderregels."],
        ["1.2", "Velden Hoeveelheid/Lengte op orderregels + native berekening "
                "(vervangt Studio-automatiseringen)."],
        ["1.1", "Gewichtsbanding per stuk voor TK1/TK2 (zwaar/midden/licht)."],
        ["1.0", "Eerste versie — transportkostenberekening voor TK1/TK2/TK3."],
    ],
)

# ============================================================
# FOOTER
# ============================================================
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("dooIT B.V. — https://dooit.nl")
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

# Save
output_path = "/home/odoo/src/user/d1_shipping_cost/static/description/Handleiding_Transportkosten.docx"
doc.save(output_path)
print(f"Handleiding opgeslagen: {output_path}")
